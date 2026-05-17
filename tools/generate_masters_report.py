from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"
OUT.mkdir(exist_ok=True)

TITLE = "EfficientDet-D0 with EA-StrongSORT for Multi-Perspective Surgical Tool Tracking on CholecTrack20"
AUTHORS = "Abdullah and Shamsa"
DATE = "May 2026"


detector_tuning_rows = [
    ["Selected", "Run", "Total epochs", "Best epoch", "AP0.5", "AP0.75", "AP0.5:0.95", "FPS"],
    ["yes", "efficientdet_d0_640_adamw_lr2e4_long", "150", "63", "44.4", "31.3", "27.9", "20.9"],
    ["no", "efficientdet_d0_512_adamw_lr2e4_batch12_workers6_eval3", "100", "96", "46.6", "28.5", "27.2", "28.8"],
    ["no", "screen_effdet_d0_640_lr2e4", "30", "30", "42.9", "29.3", "26.6", ""],
    ["no", "screen_effdet_d0_640_lr1e4", "30", "30", "41.2", "28.9", "25.7", ""],
    ["no", "screen_effdet_d0_lr1e4", "30", "25", "41.7", "26.3", "24.8", ""],
    ["no", "screen_effdet_d0_sgd_lr1e3", "30", "20", "37.3", "22.9", "21.8", ""],
]

yolo_tuning_rows = [
    ["Run", "Key settings", "Best epoch", "Precision", "Recall", "mAP50", "mAP50-95", "Decision"],
    ["final_yolo11_img768_sgd", "SGD, imgsz=768, batch=8, lr0=0.01, patience=30", "42", "0.5595", "0.4683", "0.4077", "0.2670", "Best strict localization YOLO11 run"],
    ["final_yolo11_patience30", "SGD, imgsz=640, batch=16, lr0=0.01, patience=30", "31", "0.5700", "0.4715", "0.4108", "0.2631", "Best mAP50/balanced YOLO11 run"],
    ["final_yolo11_adamw_lr001", "AdamW, imgsz=640, batch=16, lr0=0.001, patience=40", "44", "0.5573", "0.4716", "0.4035", "0.2519", "Worse than SGD"],
    ["sweep_20260503_sgd_lr001_img640", "SGD, imgsz=640, batch=16, lr0=0.01", "31", "0.5638", "0.4521", "0.4048", "0.2627", "Close to main SGD baseline"],
    ["quick_yolo_test", "SGD, imgsz=384, batch=16, epochs=10, fraction=0.35", "10", "0.3910", "0.3338", "0.2867", "0.1608", "Pipeline test only"],
]

yolo_paper_rows = [
    ["Detection model", "AP0.5", "AP0.75", "AP0.5:0.95", "Grasper", "Bipolar", "Hook", "Scissors", "Clipper", "Irrigator", "Bag", "FPS"],
    ["YOLO11-img768-SGD", "40.3", "29.6", "26.6", "42.1", "45.8", "50.0", "23.6", "59.3", "3.4", "57.8", "62.9"],
]

switch_rationale_rows = [
    ["Observation", "Evidence", "Consequence"],
    ["YOLO11 improvements saturated", "imgsz 640 to 768 improved mAP50-95 only from 0.2631 to 0.2670", "More YOLO-only tuning had diminishing returns"],
    ["AdamW did not improve YOLO11", "final_yolo11_adamw_lr001 reached mAP50-95 0.2519", "SGD remained the better YOLO11 optimizer"],
    ["Strict localization remained weak", "Best YOLO11 paper-style AP0.5:0.95 was 26.6", "Detector quality was unlikely to support strong tracking"],
    ["Irrigator remained weak", "YOLO11 irrigator AP0.5 was 3.4", "Class imbalance/appearance issues required a different detector experiment"],
    ["Supervisor direction favored EfficientNet-family study", "EfficientNet was requested as an alternative contribution", "EfficientDet-D0 was chosen as an EfficientNet-family object detector"],
]

detector_comparison_rows = [
    ["Detector", "AP0.5", "AP0.75", "AP0.5:0.95", "FPS", "Source"],
    ["YOLOv7", "80.6", "62.0", "56.1", "20.6", "CholecTrack20 GitHub"],
    ["YOLOv8", "79.1", "62.4", "55.6", "29.0", "CholecTrack20 GitHub"],
    ["YOLOv9", "80.2", "62.6", "56.5", "23.7", "CholecTrack20 GitHub"],
    ["YOLOv10", "80.1", "62.1", "55.8", "28.6", "CholecTrack20 GitHub"],
    ["FCOS", "43.5", "31.5", "28.1", "7.7", "CholecTrack20 GitHub"],
    ["YOLO11-img768-SGD", "40.3", "29.6", "26.6", "62.9", "Our baseline"],
    ["EfficientDet-D0-640-AdamW-lr2e4", "44.4", "31.2", "27.9", "20.9", "Our contribution"],
]

detector_class_rows = [
    ["Model", "Grasper", "Bipolar", "Hook", "Scissors", "Clipper", "Irrigator", "Bag"],
    ["YOLO11-img768-SGD", "42.1", "45.8", "50.0", "23.6", "59.3", "3.4", "57.8"],
    ["EfficientDet-D0-640-AdamW-lr2e4", "47.3", "50.8", "60.3", "33.4", "60.5", "0.6", "57.6"],
]

tracking_rows = [
    ["Perspective", "MOTA", "IDF1", "MOTP", "Precision", "Recall", "IDSW", "FP", "FN"],
    ["Visibility", "20.974", "24.172", "82.536", "68.692", "58.952", "3332", "8059", "12312"],
    ["Intracorporeal", "19.641", "12.158", "82.536", "68.692", "58.952", "3732", "8059", "12312"],
    ["Intraoperative", "19.314", "6.298", "82.536", "68.692", "58.952", "3830", "8059", "12312"],
]

tracking_official_rows = [
    ["Perspective", "Model", "MOTA", "IDF1", "MOTP", "FPS", "Source"],
    ["Visibility", "Bot-SORT", "72.0", "41.4", "83.7", "8.7", "CholecTrack20 GitHub"],
    ["Visibility", "ByteTrack", "69.3", "36.8", "84.0", "16.4", "CholecTrack20 GitHub"],
    ["Visibility", "SORT", "21.4", "13.4", "83.3", "19.5", "CholecTrack20 GitHub"],
    ["Visibility", "EfficientDet-D0 -> EA-StrongSORT", "21.0", "24.2", "82.5", "-", "Our contribution"],
    ["Intracorporeal", "Bot-SORT", "70.0", "18.9", "83.7", "8.7", "CholecTrack20 GitHub"],
    ["Intracorporeal", "ByteTrack", "67.4", "16.9", "84.0", "16.4", "CholecTrack20 GitHub"],
    ["Intracorporeal", "EfficientDet-D0 -> EA-StrongSORT", "19.6", "12.2", "82.5", "-", "Our contribution"],
    ["Intraoperative", "Bot-SORT", "69.6", "10.2", "83.7", "8.7", "CholecTrack20 GitHub"],
    ["Intraoperative", "ByteTrack", "67.0", "9.5", "84.0", "16.4", "CholecTrack20 GitHub"],
    ["Intraoperative", "EfficientDet-D0 -> EA-StrongSORT", "19.3", "6.3", "82.5", "-", "Our contribution"],
]

challenge_rows = [
    ["Model", "Bleeding", "Blur", "Smoke", "Crowded", "Occluded", "Foul Lens", "Trocar"],
    ["YOLO11-img768-SGD", "15.1", "100.0", "43.5", "22.8", "46.5", "12.7", "22.5"],
    ["EfficientDet-D0-640-AdamW-lr2e4", "19.2", "75.2", "46.9", "24.4", "52.0", "17.1", "47.3"],
]


sections = [
    (
        "Abstract",
        "This report presents an EfficientDet-D0 and EA-StrongSORT pipeline for multi-perspective surgical tool tracking on CholecTrack20. "
        "The study first documents a YOLO11 baseline, then replaces the detector with EfficientDet-D0, an EfficientNet-B0 plus BiFPN detector, and integrates it with EA-StrongSORT for tracking across visibility, intracorporeal, and intraoperative trajectory definitions. "
        "The selected EfficientDet-D0 detector achieved AP0.5 = 44.4, AP0.75 = 31.2, and AP0.5:0.95 = 27.9. "
        "The full tracking run produced MOTA values of 20.974, 19.641, and 19.314 for visibility, intracorporeal, and intraoperative perspectives respectively. "
        "The contribution is a reproducible EfficientNet-family detector-to-tracker baseline with complete metric logging and analysis, rather than a claim of state-of-the-art performance.",
    ),
    (
        "1. Introduction",
        "Surgical tool tracking in laparoscopic video is a difficult multi-object tracking problem because instruments frequently leave the camera view, are partially occluded, overlap, and appear under smoke, blur, blood, reflection, and lens fouling. "
        "CholecTrack20 addresses this by providing multi-perspective annotations for surgical instruments, including visibility, intracorporeal, and intraoperative track identities. "
        "The goal of this project was to extend the CholecTrack20 research direction with an EfficientNet-family detection front end followed by EA-StrongSORT tracking.",
    ),
    (
        "2. Contribution",
        "The implemented contribution has four parts. First, the project establishes a documented YOLO11 baseline. Second, it introduces EfficientDet-D0 as the active detector, selected because it is an EfficientNet-family object detector and is feasible on an RTX 2060 GPU. "
        "Third, it connects EfficientDet-D0 detections directly to EA-StrongSORT without relying on BoxMOT detector-name routing. Fourth, it adds tracking evaluation against CholecTrack20 JSON track IDs and exports reproducible CSV and Markdown logs for detector and tracker experiments.",
    ),
    (
        "3. Dataset and Evaluation Context",
        "CholecTrack20 contains 20 laparoscopic cholecystectomy videos split into 10 training, 2 validation, and 8 testing videos. "
        "The dataset provides seven tool classes: grasper, bipolar, hook, scissors, clipper, irrigator, and specimen bag. "
        "For tracking, it provides three identity definitions: visibility trajectory, intracorporeal trajectory, and intraoperative trajectory. "
        "The public repository reports detector AP metrics and tracking metrics including MOTA, MOTP, IDF1, and HOTA-style measures. "
        "Our report matches the public detector and tracker columns where possible, while noting that the tracking evaluator implemented here is an internal class-aware IoU evaluator and not an official TrackEval leaderboard submission.",
    ),
    (
        "4. YOLO11 Baseline and Switch Rationale",
        "The project began with YOLO11 because it is a modern one-stage detector and gave a fast baseline for the detection-to-tracking pipeline. "
        "Several YOLO11 configurations were tested, including SGD at image sizes 640 and 768, AdamW at image size 640, and quick partial-dataset runs for pipeline validation. "
        "The best YOLO11 strict localization run was final_yolo11_img768_sgd, with mAP50-95 = 0.2670 in the training logs and paper-style AP0.5:0.95 = 26.6. "
        "Although increasing image size from 640 to 768 gave a small improvement, the gain was not large enough to justify continuing only with YOLO11. "
        "The doctor/supervisor feedback also requested an EfficientNet-family alternative. Since plain EfficientNet is a classifier, the project selected EfficientDet-D0, which uses an EfficientNet-B0 backbone and is a real object detector. "
        "Therefore, YOLO11 was kept as the documented baseline, and EfficientDet-D0 became the active detector for the final EA-StrongSORT contribution.",
    ),
    (
        "5. Method",
        "The final detector is EfficientDet-D0 using a COCO-pretrained EfficientNet-B0 backbone and BiFPN detection head. "
        "The selected training run used image size 640, AdamW optimization, learning rate 0.0002, weight decay 0.0001, batch size 6, gradient accumulation of 2, six workers, 150 total epochs, and validation every three epochs. "
        "The best checkpoint occurred at epoch 63, after which validation performance declined slightly. "
        "For tracking, EfficientDet detections were converted to [x1, y1, x2, y2, confidence, class] format and passed into StrongSORT with OSNet ReID weights. "
        "Predictions were exported in MOT text format and evaluated against the three CholecTrack20 perspective-specific track ID fields.",
    ),
    (
        "6. Detector Experiments",
        "Detector tuning used short screening runs before longer training. The 150-epoch EfficientDet-D0 run at image size 640 and learning rate 0.0002 produced the best AP0.5:0.95 and was selected for tracking. "
        "It improved over our YOLO11 baseline in AP0.5, AP0.75, and AP0.5:0.95, but it remained below the strongest public CholecTrack20 YOLOv7-v10 detector benchmarks.",
    ),
    (
        "7. Tracking Experiments",
        "The selected EfficientDet-D0 checkpoint was evaluated with EA-StrongSORT on all eight CholecTrack20 test videos. "
        "The benchmark produced 768,555 MOT-format prediction rows across the test set. "
        "Tracking evaluation used annotated CholecTrack20 frames only with class-aware IoU >= 0.5 matching. "
        "MOTP remained high at 82.536, showing that matched boxes were localized reasonably well, but MOTA and IDF1 were limited by missed detections, false positives, and identity fragmentation.",
    ),
    (
        "8. Comparison with the CholecTrack20 GitHub Benchmark",
        "Against the public detection leaderboard, EfficientDet-D0 is comparable to older detector families such as FCOS on AP0.5:0.95 but far behind the strongest YOLOv7-v10 benchmarks. "
        "Against the public tracking leaderboard, our EA-StrongSORT pipeline does not outperform Bot-SORT or ByteTrack. "
        "However, it provides a new EfficientNet-family detector-to-tracker baseline and shows that detector quality is the main bottleneck before tracker parameter optimization can be expected to help substantially.",
    ),
    (
        "9. Limitations",
        "The main limitation is detector quality. The selected EfficientDet-D0 model reaches AP0.5:0.95 = 27.9, whereas the strongest public CholecTrack20 detectors exceed 55 on the same AP column. "
        "The irrigator class remains especially weak, suggesting class imbalance and visual ambiguity. "
        "The tracker also produces many identity switches, especially under intracorporeal and intraoperative identity definitions, where out-of-view and out-of-body identity continuity is stricter than frame-visible tracking.",
    ),
    (
        "10. Future Work",
        "Future work should prioritize improving detection through stronger EfficientDet variants if GPU memory allows, class rebalancing, hard-example mining, and additional data augmentation for rare tools. "
        "After detector quality improves, EA-StrongSORT can be tuned through confidence thresholds, max age, ReID distance thresholds, and motion compensation behavior. "
        "A final official comparison should use the original CholecTrack20 TrackEval adaptation so that HOTA and leaderboard values are directly comparable.",
    ),
    (
        "11. Conclusion",
        "The project successfully implements and documents an EfficientDet-D0 to EA-StrongSORT pipeline for CholecTrack20. "
        "While it does not exceed the strongest public benchmark results, it adds a reproducible EfficientNet-family baseline, demonstrates end-to-end multi-perspective tracking, and identifies detector quality and identity fragmentation as the key barriers to stronger performance.",
    ),
]


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in str(text))


def latex_table(caption: str, label: str, rows: list[list[str]]) -> str:
    cols = "l" * len(rows[0])
    lines = [
        r"\begin{table}[H]",
        r"\centering",
        r"\small",
        r"\resizebox{\textwidth}{!}{%",
        rf"\begin{{tabular}}{{{cols}}}",
        r"\toprule",
        " & ".join(latex_escape(cell) for cell in rows[0]) + r" \\",
        r"\midrule",
    ]
    for row in rows[1:]:
        lines.append(" & ".join(latex_escape(cell) for cell in row) + r" \\")
    lines.extend(
        [
            r"\bottomrule",
            r"\end{tabular}%",
            r"}",
            rf"\caption{{{latex_escape(caption)}}}",
            rf"\label{{{label}}}",
            r"\end{table}",
        ]
    )
    return "\n".join(lines)


def write_latex() -> None:
    bib = dedent(
        r"""
        @inproceedings{nwoye2025cholectrack20,
          author    = {Nwoye, Chinedu Innocent and Elgohary, Kareem and Srinivas, Anvita and Zaid, Fauzan and Lavanchy, Joel L. and Padoy, Nicolas},
          title     = {CholecTrack20: A Multi-Perspective Tracking Dataset for Surgical Tools},
          booktitle = {Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition},
          year      = {2025}
        }

        @misc{cholectrack20github,
          author       = {{CAMMA}},
          title        = {CholecTrack20 GitHub Repository and Benchmark Tables},
          year         = {2025},
          howpublished = {\url{https://github.com/CAMMA-public/cholectrack20}},
          note         = {Accessed 12 May 2026}
        }

        @misc{efficientdet,
          author       = {Tan, Mingxing and Pang, Ruoming and Le, Quoc V.},
          title        = {EfficientDet: Scalable and Efficient Object Detection},
          year         = {2020},
          eprint       = {1911.09070},
          archivePrefix= {arXiv}
        }

        @inproceedings{strongsort,
          author    = {Du, Yunhao and Song, Yang and Yang, Bo and Zhao, Yanyun},
          title     = {StrongSORT: Make DeepSORT Great Again},
          booktitle = {IEEE Transactions on Multimedia},
          year      = {2023}
        }
        """
    ).strip()
    (OUT / "references.bib").write_text(bib + "\n", encoding="utf-8")

    body_parts = []
    for title, text in sections:
        if title == "Abstract":
            body_parts.append(r"\begin{abstract}" + "\n" + latex_escape(text) + "\n" + r"\end{abstract}")
        else:
            body_parts.append(r"\section{" + latex_escape(title.split(". ", 1)[-1]) + "}")
            body_parts.append(latex_escape(text))
            if title.startswith("4."):
                body_parts.append(latex_table("YOLO11 hyperparameter experiments before switching detector family.", "tab:yolo_tuning", yolo_tuning_rows))
                body_parts.append(latex_table("Best YOLO11 paper-style detector row.", "tab:yolo_paper", yolo_paper_rows))
                body_parts.append(latex_table("Rationale for switching from YOLO11 to EfficientDet-D0.", "tab:switch_rationale", switch_rationale_rows))
            if title.startswith("6."):
                body_parts.append(latex_table("EfficientDet-D0 hyperparameter screening and selected detector.", "tab:detector_tuning", detector_tuning_rows))
                body_parts.append(latex_table("Detector comparison using CholecTrack20-style AP columns.", "tab:detector_compare", detector_comparison_rows))
                body_parts.append(latex_table("Per-class AP at IoU threshold 0.5 for internal detector comparison.", "tab:class_ap", detector_class_rows))
                body_parts.append(latex_table("Detection AP across visual challenge subsets.", "tab:challenge_ap", challenge_rows))
            if title.startswith("7."):
                body_parts.append(latex_table("EA-StrongSORT tracking metrics from the implemented evaluator.", "tab:tracking_ours", tracking_rows))
            if title.startswith("8."):
                body_parts.append(latex_table("Tracking comparison against selected CholecTrack20 GitHub benchmark entries.", "tab:tracking_compare", tracking_official_rows))

    tex = dedent(
        rf"""
        \documentclass[12pt,a4paper]{{report}}
        \usepackage[margin=1in]{{geometry}}
        \usepackage{{booktabs}}
        \usepackage{{array}}
        \usepackage{{graphicx}}
        \usepackage{{float}}
        \usepackage{{hyperref}}
        \usepackage{{longtable}}
        \usepackage{{caption}}
        \usepackage{{setspace}}
        \usepackage{{titlesec}}
        \onehalfspacing
        \hypersetup{{colorlinks=true, linkcolor=blue, citecolor=blue, urlcolor=blue}}
        \title{{{latex_escape(TITLE)}}}
        \author{{{latex_escape(AUTHORS)}}}
        \date{{{latex_escape(DATE)}}}
        \begin{{document}}
        \maketitle
        \tableofcontents
        \clearpage
        {chr(10).join(body_parts)}
        \clearpage
        \bibliographystyle{{IEEEtran}}
        \bibliography{{references}}
        \end{{document}}
        """
    ).strip()
    (OUT / "efficientdet_ea_strongsort_report.tex").write_text(tex + "\n", encoding="utf-8")


def pdf_table(rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = Table(rows, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F6FA")]),
            ]
        )
    )
    return table


def write_pdf() -> None:
    path = OUT / "efficientdet_ea_strongsort_report.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontSize=20, leading=24))
    styles.add(ParagraphStyle(name="BodyJustify", parent=styles["BodyText"], alignment=TA_JUSTIFY, fontSize=9.5, leading=13))
    styles.add(ParagraphStyle(name="Heading", parent=styles["Heading1"], alignment=TA_LEFT, fontSize=14, leading=18, spaceBefore=10, spaceAfter=6))
    story = [
        Paragraph(TITLE, styles["TitleCenter"]),
        Spacer(1, 0.4 * cm),
        Paragraph(AUTHORS, styles["TitleCenter"]),
        Paragraph(DATE, styles["BodyText"]),
        Spacer(1, 0.8 * cm),
    ]
    for title, text in sections:
        story.append(Paragraph(title, styles["Heading"]))
        story.append(Paragraph(text, styles["BodyJustify"]))
        story.append(Spacer(1, 0.25 * cm))
        if title.startswith("4."):
            story.append(Paragraph("YOLO11 hyperparameter experiments", styles["Heading"]))
            story.append(pdf_table(yolo_tuning_rows))
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph("Best YOLO11 paper-style detector row", styles["Heading"]))
            story.append(pdf_table(yolo_paper_rows))
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph("Why the detector was switched to EfficientDet-D0", styles["Heading"]))
            story.append(pdf_table(switch_rationale_rows))
        if title.startswith("6."):
            story.append(Paragraph("Detector hyperparameter screening", styles["Heading"]))
            story.append(pdf_table(detector_tuning_rows))
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph("Detector comparison with CholecTrack20-style AP columns", styles["Heading"]))
            story.append(pdf_table(detector_comparison_rows))
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph("Per-class detector comparison", styles["Heading"]))
            story.append(pdf_table(detector_class_rows))
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph("Visual challenge AP comparison", styles["Heading"]))
            story.append(pdf_table(challenge_rows))
        if title.startswith("7."):
            story.append(Paragraph("Implemented EA-StrongSORT tracking metrics", styles["Heading"]))
            story.append(pdf_table(tracking_rows))
        if title.startswith("8."):
            story.append(Paragraph("Comparison with selected CholecTrack20 GitHub tracking entries", styles["Heading"]))
            story.append(pdf_table(tracking_official_rows))
    story.append(PageBreak())
    story.append(Paragraph("References", styles["Heading"]))
    refs = [
        "Nwoye et al., CholecTrack20: A Multi-Perspective Tracking Dataset for Surgical Tools, CVPR 2025.",
        "CAMMA-public CholecTrack20 GitHub repository and benchmark tables, accessed 12 May 2026, https://github.com/CAMMA-public/cholectrack20.",
        "Tan et al., EfficientDet: Scalable and Efficient Object Detection, 2020.",
        "Du et al., StrongSORT: Make DeepSORT Great Again, IEEE Transactions on Multimedia, 2023.",
    ]
    for ref in refs:
        story.append(Paragraph(ref, styles["BodyJustify"]))
    doc.build(story)


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_paragraph()


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(18)
    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.add_run(f"{AUTHORS}\n{DATE}").font.size = Pt(11)

    for title_text, text in sections:
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(text)
        if title_text.startswith("4."):
            doc.add_heading("YOLO11 hyperparameter experiments", level=2)
            add_docx_table(doc, yolo_tuning_rows)
            doc.add_heading("Best YOLO11 paper-style detector row", level=2)
            add_docx_table(doc, yolo_paper_rows)
            doc.add_heading("Why the detector was switched to EfficientDet-D0", level=2)
            add_docx_table(doc, switch_rationale_rows)
        if title_text.startswith("6."):
            doc.add_heading("Detector hyperparameter screening", level=2)
            add_docx_table(doc, detector_tuning_rows)
            doc.add_heading("Detector comparison with CholecTrack20-style AP columns", level=2)
            add_docx_table(doc, detector_comparison_rows)
            doc.add_heading("Per-class detector comparison", level=2)
            add_docx_table(doc, detector_class_rows)
            doc.add_heading("Visual challenge AP comparison", level=2)
            add_docx_table(doc, challenge_rows)
        if title_text.startswith("7."):
            doc.add_heading("Implemented EA-StrongSORT tracking metrics", level=2)
            add_docx_table(doc, tracking_rows)
        if title_text.startswith("8."):
            doc.add_heading("Comparison with selected CholecTrack20 GitHub tracking entries", level=2)
            add_docx_table(doc, tracking_official_rows)
    doc.add_heading("References", level=1)
    for ref in [
        "Nwoye et al., CholecTrack20: A Multi-Perspective Tracking Dataset for Surgical Tools, CVPR 2025.",
        "CAMMA-public CholecTrack20 GitHub repository and benchmark tables, accessed 12 May 2026, https://github.com/CAMMA-public/cholectrack20.",
        "Tan et al., EfficientDet: Scalable and Efficient Object Detection, 2020.",
        "Du et al., StrongSORT: Make DeepSORT Great Again, IEEE Transactions on Multimedia, 2023.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")
    doc.save(OUT / "efficientdet_ea_strongsort_report.docx")


def main() -> None:
    write_latex()
    write_pdf()
    write_docx()
    print(f"Wrote report files to {OUT}")


if __name__ == "__main__":
    main()
